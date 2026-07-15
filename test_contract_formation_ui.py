"""Test UI reel + non-regression - Sprint 18.1, Contrat de cession <-> Formations.

Script autonome (meme convention que les tests precedents), base SQLite
TEMPORAIRE et isolee, vraie QApplication PySide6 offscreen, QMessageBox
patchee (voir test_cddu_ui.py pour l'explication du blocage headless).

Verifie : creation d'un contrat depuis une prestation en un clic (aucune
selection d'artiste requise), prerempli integralement depuis la Formation,
compatibilite d'un ancien contrat (artist_id, sans formation_id), generation
DOCX/PDF reelle affichant le nom de la Formation et JAMAIS un nom de
musicien individuel, et absence totale de reference aux
participants/musiciens sur le contrat de cession."""

from __future__ import annotations

import os
import tempfile
from pathlib import Path
from unittest.mock import patch

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")

import app.database.database as database_module

_tmp_dir = tempfile.TemporaryDirectory()
database_module.DB_PATH = Path(_tmp_dir.name) / "test_contract_formation_ui.db"
database_module.Database._instance = None

from PySide6.QtWidgets import QApplication, QMessageBox  # noqa: E402

app = QApplication.instance() or QApplication([])

_info_patch = patch.object(QMessageBox, "information", return_value=QMessageBox.StandardButton.Ok)
_warning_patch = patch.object(QMessageBox, "warning", return_value=QMessageBox.StandardButton.Ok)
_question_patch = patch.object(QMessageBox, "question", return_value=QMessageBox.StandardButton.Yes)
_info_patch.start()
_warning_patch.start()
_question_patch.start()

from app.database.migrations import MigrationManager  # noqa: E402
from app.models.artist import Artist  # noqa: E402
from app.models.contract import Contract  # noqa: E402
from app.models.formation import Formation  # noqa: E402
from app.models.organization import Organization  # noqa: E402
from app.models.prestation import Prestation  # noqa: E402
from app.models.producteur import Producteur  # noqa: E402
from app.services.artist_service import ArtistService  # noqa: E402
from app.services.contract_service import ContractService  # noqa: E402
from app.services.formation_artiste_service import FormationArtisteService  # noqa: E402
from app.services.formation_service import FormationService  # noqa: E402
from app.services.organization_service import OrganizationService  # noqa: E402
from app.services.prestation_service import PrestationService  # noqa: E402
from app.services.producteur_service import ProducteurService  # noqa: E402
from app.ui.contract_dialog import ContractDialog  # noqa: E402

failures: list[str] = []


def check(label: str, condition: bool) -> None:
    status = "OK" if condition else "ECHEC"
    print(f"[{status}] {label}")
    if not condition:
        failures.append(label)


MigrationManager().migrate()

# ---------------------------------------------------------------------------
print("== Donnees de test ==")

producteur_service = ProducteurService()
producteur_service.create_producteur(Producteur(
    nom="YGNT PRODUCTION", forme_juridique="Association loi 1901",
    city="Villeneuve", siret="10572314200011", representant="Tanguy Zahn", fonction="President",
))

formation_service = FormationService()
composition_service = FormationArtisteService()
formation_id = formation_service.create_formation(Formation(
    nom="SANFUEGO", style="Flamenco",
    address="12 rue des Artistes", postal_code="83000", city="Toulon",
    phone="0400000000", email="contact@sanfuego.example",
    siret="98765432100012", ape="9001Z", licence="PLATESV-R-2026-000001",
    iban="FR7630006000011234567890189", bic="AGRIFRPP",
))

artist_service = ArtistService()
anthony_id = artist_service.create_artist(Artist(legal_name="Anthony Zahn", instrument="Guitariste", qualification="Artiste musicien"))
miguel_id = artist_service.create_artist(Artist(legal_name="Miguel", instrument="Batteur", qualification="Artiste musicien"))
mahoni_id = artist_service.create_artist(Artist(legal_name="Mahoni", instrument="Chanteur", qualification="Chanteur"))
for artist_id, role in ((anthony_id, "Guitare"), (miguel_id, "Batterie"), (mahoni_id, "Chant")):
    composition_service.add_member(formation_id, artist_id, role=role)

organization_service = OrganizationService()
organization_id = organization_service.create_organization(Organization(
    name="Mairie de Testville", address="1 place de la Mairie", city="Testville", siret="11122233300099",
))

prestation_service = PrestationService()
prestation_id = prestation_service.create_prestation(Prestation(
    nom="Festival Ete 2026", type_evenement="festival", date_debut="09/07/2026",
    formation_id=formation_id, organization_id=organization_id,
    lieu_nom="Plaine des fetes", lieu_city="Testville",
))
prestation = prestation_service.get_prestation(prestation_id)

# ---------------------------------------------------------------------------
print("\n== Creation depuis une prestation : un clic, aucune selection d'artiste ==")

contract_service = ContractService(formation_service=formation_service)

# Reproduit exactement PrestationsPage.create_contract_from_selected_prestation() :
# un seul appel, sans aucune autre intervention.
seed = contract_service.build_from_prestation(prestation)

check("formation_id lu automatiquement depuis la prestation", seed.formation_id == formation_id)
check("artist_id (legacy) reste None : jamais impose", seed.artist_id is None)
check("organisateur pre-rempli automatiquement (organization_id)", seed.organization_id == organization_id)
check("lieu pre-rempli automatiquement", seed.prestation_lieu == "Plaine des fetes")
check("date pre-remplie automatiquement", seed.prestation_date == "09/07/2026")
check("nom de la Formation pre-rempli sans aucune action utilisateur", seed.artiste_nom == "SANFUEGO")
check("adresse de la Formation pre-remplie", seed.artiste_adresse == "12 rue des Artistes")
check("SIRET de la Formation pre-rempli", seed.artiste_siret == "98765432100012")
check("IBAN de la Formation pre-rempli", seed.artiste_iban == "FR7630006000011234567890189")
check(
    "spectacle_nom pre-rempli avec le nom de la Formation (seul champ reellement imprime par le template)",
    seed.spectacle_nom == "SANFUEGO",
)

dialog = ContractDialog(
    None,
    initial_contract=seed,
    service=contract_service,
    formation_service=formation_service,
    organization_service=organization_service,
)
dialog.show()

check("onglet renomme 'Formation' (plus 'Artiste')", "Formation" in [dialog.tabs.tabText(i) for i in range(dialog.tabs.count())])
check("aucun onglet 'Artiste' residuel", "Artiste" not in [dialog.tabs.tabText(i) for i in range(dialog.tabs.count())])
check("le combo Formation est preselectionne sans action de l'utilisateur", dialog.formation_combo.currentData() == formation_id)
check("le champ Nom affiche deja 'SANFUEGO' a l'ouverture du dialogue", dialog.artiste_nom.text() == "SANFUEGO")
check("le champ Organisateur affiche deja 'Mairie de Testville' a l'ouverture", dialog.organisateur.text() == "Mairie de Testville")
check("le champ Spectacle affiche deja 'SANFUEGO' a l'ouverture (aucune saisie requise)", dialog.spectacle.text() == "SANFUEGO")

# Aucune saisie manuelle : le contrat est deja complet et enregistrable en
# un clic, exactement comme demande ("Un clic. -> Contrat prerempli.").
dialog.save()

check("le contrat est enregistre", dialog._source_contract is not None and dialog._source_contract.id is not None)
contract_id = dialog._source_contract.id

saved = contract_service.get_contract(contract_id)
check("formation_id persiste en base", saved.formation_id == formation_id)
check("artist_id reste None (jamais ecrit par ce dialogue)", saved.artist_id is None)
check("instantane artiste_nom = nom de la Formation", saved.artiste_nom == "SANFUEGO")

# ---------------------------------------------------------------------------
print("\n== Generation DOCX reelle : Formation oui, musiciens jamais ==")

docx_path = contract_service.generate_docx(contract_id)
check("le fichier DOCX existe", docx_path.exists())

from docx import Document  # noqa: E402

doc = Document(str(docx_path))
full_text = "\n".join(p.text for p in doc.paragraphs)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            full_text += "\n" + "\n".join(p.text for p in cell.paragraphs)

check("aucun placeholder {{...}} residuel", "{{" not in full_text and "}}" not in full_text)
check(
    "le nom de la Formation apparait dans le DOCX (via spectacle_nom, pre-rempli automatiquement)",
    "SANFUEGO" in full_text and saved.spectacle_nom == "SANFUEGO",
)
check("Anthony n'apparait jamais dans le DOCX", "Anthony" not in full_text)
check("Miguel n'apparait jamais dans le DOCX", "Miguel" not in full_text)
check("Mahoni n'apparait jamais dans le DOCX", "Mahoni" not in full_text)
check("le nombre de musiciens n'apparait nulle part (aucun effectif imprime)", "3 musiciens" not in full_text)

# ---------------------------------------------------------------------------
print("\n== Generation PDF reelle (necessite Microsoft Word) ==")

try:
    pdf_path = contract_service.export_pdf(contract_id)
    check("le fichier PDF existe", pdf_path.exists())

    import fitz  # noqa: E402

    pdf_doc = fitz.open(str(pdf_path))
    pdf_text = "".join(page.get_text() for page in pdf_doc)
    check("le PDF affiche le nom de la Formation", "SANFUEGO" in pdf_text)
    check("le PDF n'affiche aucun nom de musicien", all(name not in pdf_text for name in ("Anthony", "Miguel", "Mahoni")))
except Exception as exc:
    print(f"[INFO] Conversion PDF non verifiee sur ce poste : {exc!r}")

# ---------------------------------------------------------------------------
print("\n== Compatibilite : ancien contrat (artist_id, sans formation_id) ==")

legacy_contract = Contract(
    contract_number=contract_service.next_contract_number(),
    artist_id=anthony_id,
    formation_id=None,
    organisateur_structure="Ancien Organisateur",
    spectacle_nom="Ancien spectacle",
    artiste_nom="Anthony Zahn (ancienne fiche Artiste)",
    artiste_adresse="Ancienne adresse enregistree",
    artiste_siret="00000000000000",
)
legacy_id = contract_service.create_contract(legacy_contract)
legacy_saved = contract_service.get_contract(legacy_id)

check("l'ancien contrat garde son artist_id, sans formation_id", legacy_saved.artist_id == anthony_id and legacy_saved.formation_id is None)

legacy_dialog = ContractDialog(None, contract=legacy_saved, service=contract_service, formation_service=formation_service, organization_service=organization_service)
legacy_dialog.show()

check("le combo Formation reste sur '(Aucune)' pour un ancien contrat", legacy_dialog.formation_combo.currentData() is None)
check(
    "les donnees figees de l'ancien contrat restent affichees telles quelles",
    legacy_dialog.artiste_nom.text() == "Anthony Zahn (ancienne fiche Artiste)"
    and legacy_dialog.artiste_adresse.text() == "Ancienne adresse enregistree",
)

# Regeneration DOCX d'un ancien contrat : ne doit jamais casser (Regle n°1).
# Note : le template contrat_cession.docx n'a jamais imprime artiste_nom
# (verifie plus haut - aucun placeholder {{artiste_*}} n'existe dans le
# fichier), donc la regeneration n'est pas censee faire apparaitre cette
# valeur a l'impression ; ce qui compte est qu'elle reste intacte en base
# et que la regeneration elle-meme ne leve aucune erreur (non-regression).
legacy_docx_path = contract_service.generate_docx(legacy_id)
check("un ancien contrat peut toujours regenerer son DOCX sans erreur", legacy_docx_path.exists())
legacy_reloaded = contract_service.get_contract(legacy_id)
check(
    "la regeneration DOCX ne corrompt pas les donnees figees d'origine de l'ancien contrat",
    legacy_reloaded.artiste_nom == "Anthony Zahn (ancienne fiche Artiste)"
    and legacy_reloaded.artiste_adresse == "Ancienne adresse enregistree",
)

# ---------------------------------------------------------------------------
print("\n== Non-regression : aucune reference a prestation_participants/CDDU ==")

import inspect  # noqa: E402

from app.services import contract_service as contract_service_module  # noqa: E402

source = inspect.getsource(contract_service_module)
check("ContractService ne mentionne jamais prestation_participants", "prestation_participants" not in source and "PrestationParticipant" not in source)

from app.ui import contract_dialog as contract_dialog_module  # noqa: E402

# Recherche un usage reel du code (import du service, appel de methode),
# pas une simple presence textuelle : le fichier contient volontairement un
# commentaire expliquant que prestation_participants n'est JAMAIS utilise
# ici, ce qui ferait echouer une recherche textuelle naive.
dialog_source = inspect.getsource(contract_dialog_module)
check(
    "ContractDialog n'importe ni n'utilise jamais PrestationParticipantService",
    "PrestationParticipantService" not in dialog_source
    and ".list_participants(" not in dialog_source
    and ".add_participant(" not in dialog_source,
)
check("ContractDialog n'importe jamais FormationArtisteService (composition, reservee a la fiche Formation)", "FormationArtisteService" not in dialog_source)

# ---------------------------------------------------------------------------
print("\n== Nettoyage ==")

_info_patch.stop()
_warning_patch.stop()
_question_patch.stop()

database_module.Database.instance().close()
_tmp_dir.cleanup()
print("Base temporaire supprimee.")

print("\n" + "=" * 60)
if failures:
    print(f"{len(failures)} verification(s) en echec :")
    for label in failures:
        print(f"  - {label}")
    raise SystemExit(1)

print("Toutes les verifications sont passees.")
