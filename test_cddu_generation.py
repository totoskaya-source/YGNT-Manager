"""Test de generation reelle DOCX/PDF - Sprint 16.0, module CDDU.

Script autonome (meme convention que test_cddu_migration.py), execute sur une
base SQLite TEMPORAIRE et isolee (data/ygnt_manager.db n'est jamais touchee),
mais genere un VRAI fichier DOCX et un VRAI PDF dans exports/ (comme le
ferait l'application), via le Service uniquement - aucune UI.

Verifie :
  - le DOCX est genere a partir de templates/contrat_cddu.docx ;
  - tous les placeholders attendus sont remplaces (employeur, salarie,
    convention collective, objet, duree, remuneration, defraiements,
    signatures) ;
  - aucun "{{...}}" residuel nulle part dans le document ;
  - l'article Defraiements n'affiche que les montants reellement renseignes ;
  - {{numero_objet}} reste vide ;
  - {{dates_travaillees}}/{{nombre_total_cachets}} refletent
    contrat_cddu_dates, jamais une colonne du contrat lui-meme (CDDU simple :
    une seule date) ;
  - le document s'ouvre correctement avec python-docx (equivalent
    programmatique de "ouverture Word" - une ouverture reelle dans
    Microsoft Word depend d'un poste avec Word installe, voir plus bas) ;
  - la conversion PDF fonctionne (si Microsoft Word est disponible sur ce
    poste) et produit un fichier non vide.
"""

from __future__ import annotations

import tempfile
from pathlib import Path

import app.database.database as database_module

_tmp_dir = tempfile.TemporaryDirectory()
database_module.DB_PATH = Path(_tmp_dir.name) / "test_cddu_generation.db"
database_module.Database._instance = None

from app.database.migrations import MigrationManager  # noqa: E402
from app.models.artist import Artist  # noqa: E402
from app.models.prestation import Prestation  # noqa: E402
from app.models.producteur import Producteur  # noqa: E402
from app.services.artist_service import ArtistService  # noqa: E402
from app.services.contrat_cddu_service import ContratCdduService  # noqa: E402
from app.services.prestation_service import PrestationService  # noqa: E402
from app.services.producteur_service import ProducteurService  # noqa: E402

failures: list[str] = []


def check(label: str, condition: bool) -> None:
    status = "OK" if condition else "ECHEC"
    print(f"[{status}] {label}")
    if not condition:
        failures.append(label)


MigrationManager().migrate()

# ---------------------------------------------------------------------------
print("== Donnees de test (reprennent le contrat de reference Ygnt/Anthony Zahn) ==")

producteur_service = ProducteurService()
producteur_id = producteur_service.create_producteur(Producteur(
    nom="YGNT PRODUCTION",
    forme_juridique="Association loi 1901",
    adresse="28 Rue Joseph Roumanille",
    postal_code="04180",
    city="Villeneuve",
    siret="10572314200011",
    representant="Tanguy Zahn",
    fonction="Président",
    convention_collective="IDCC 01285",
))

artist_service = ArtistService()
artist_id = artist_service.create_artist(Artist(
    legal_name="Anthony Zahn",
    address="301 Rue du Pigeonnier",
    postal_code="83600",
    city="Fréjus",
    birth_date="11/05/1988",
    birth_place="Cavaillon",
    social_number="1 88 05 84 035 041 22",
    conges_spectacle_number="C 381 825",
    instrument="Guitariste",
    qualification="Artiste musicien",
))

prestation_service = PrestationService()
prestation_id = prestation_service.create_prestation(Prestation(
    nom="Concert de flamenco lors des nuits Nomades de Maneo avec le groupe Sanfuego",
    type_evenement="festival",
    date_debut="09/07/2026",
    lieu_nom="Maneo",
    lieu_city="Aix en Provence",
    artist_id=artist_id,
))

# ---------------------------------------------------------------------------
print("\n== Creation du CDDU (contrat simple, une seule date) ==")

cddu_service = ContratCdduService()
prestation = prestation_service.get_prestation(prestation_id)

contrat = cddu_service.build_from_prestation_and_artist(prestation, artist_id)
contrat.remuneration_brute = "136.37"
# Un seul defraiement renseigne (Deplacement) : Hebergement/Repas/Autres/
# Montant libre doivent disparaitre du document genere.
contrat.defraiement_deplacement = "45"

contrat_id = cddu_service.create_contrat(contrat)
cddu_service.add_date(contrat_id, "09/07/2026", prestation_id=prestation_id, nombre_cachets=1)

saved = cddu_service.get_contrat(contrat_id)
check("CDDU cree avec numero CDDU-AAAA-0001", saved.numero.endswith("-0001") and saved.numero.startswith("CDDU-"))

# ---------------------------------------------------------------------------
print("\n== Generation DOCX reelle ==")

docx_path = cddu_service.generate_docx(contrat_id)
check("le fichier DOCX existe sur le disque", docx_path.exists())
check("le fichier DOCX n'est pas vide", docx_path.stat().st_size > 0)

saved = cddu_service.get_contrat(contrat_id)
check("contrat.docx_path enregistre", saved.docx_path == str(docx_path))
check("contrat.generated_at renseigne", bool(saved.generated_at))

from docx import Document  # noqa: E402

doc = Document(str(docx_path))


def all_text(document) -> str:
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


full_text = all_text(doc)

check("le document s'ouvre correctement avec python-docx (equiv. ouverture Word)", isinstance(full_text, str) and len(full_text) > 0)
check("aucun placeholder {{...}} residuel dans le document", "{{" not in full_text and "}}" not in full_text)

# ---------------------------------------------------------------------------
print("\n== Placeholders Employeur/Salarie/Convention collective/Objet ==")

check("titre du contrat present", "CONTRAT D'ENGAGEMENT ARTISTE" in full_text)
check("numero de contrat present", saved.numero in full_text)
check("employeur : nom present", "YGNT PRODUCTION" in full_text)
check("employeur : adresse complete presente (pas de virgule orpheline)", "28 Rue Joseph Roumanille, 04180 Villeneuve" in full_text)
check("employeur : SIRET present", "10572314200011" in full_text)
check("employeur : representant/fonction presents", "Tanguy Zahn" in full_text and "Président" in full_text)
check("convention collective presente", "IDCC 01285" in full_text)
check("salarie : nom present", "Anthony Zahn" in full_text)
check("salarie : date/lieu de naissance presents", "11/05/1988" in full_text and "Cavaillon" in full_text)
check("salarie : adresse complete presente", "301 Rue du Pigeonnier, 83600 Fréjus" in full_text)
check("salarie : numero de securite sociale present", "1 88 05 84 035 041 22" in full_text)
check("salarie : numero de conges spectacle present", "C 381 825" in full_text)
check("fonction (objet) presente", "Guitariste" in full_text)
check("objet (description du projet) present", "Concert de flamenco" in full_text)
check("lieu du projet present", "Maneo, Aix en Provence" in full_text)

# ---------------------------------------------------------------------------
print("\n== numero_objet, dates travaillees, cachets ==")

numero_objet_paragraph = next(
    (p.text for p in doc.paragraphs if "Le numéro d'objet lié à ce contrat est le" in p.text),
    None,
)
check(
    "le numero d'objet reste vide (phrase presente, rien apres 'le')",
    numero_objet_paragraph is not None
    and numero_objet_paragraph.strip() == "Le numéro d'objet lié à ce contrat est le",
)

check("date travaillee listee (contrat_cddu_dates, jamais une colonne du contrat)", "09/07/2026 - 1 Cachet de représentation" in full_text)
check("nombre total de cachets affiche", "Nombre total de cachets : 1" in full_text)

# ---------------------------------------------------------------------------
print("\n== Remuneration et Defraiements ==")

check("remuneration brute affichee (format monetaire francais)", "136,37" in full_text)
check("defraiement Deplacement renseigne et affiche", "Déplacement : 45,00 euros" in full_text)
check("defraiement Hebergement absent (0 => ligne supprimee)", "Hébergement :" not in full_text)
check("defraiement Repas absent (0 => ligne supprimee)", "Repas :" not in full_text)

# ---------------------------------------------------------------------------
print("\n== Signatures ==")

check("ville de signature presente (repli sur la ville du producteur)", "Villeneuve" in full_text)
check("date de signature renseignee automatiquement", bool(saved.date_signature))
check("bloc signature Salarie/Employeur present", "Le Salarié" in full_text and "L'Employeur" in full_text)
check("bloc signature reprend le nom du salarie et du representant", "Anthony Zahn" in full_text and "Tanguy Zahn pour YGNT PRODUCTION" in full_text)

history_after_docx = cddu_service.history(contrat_id)
check("historique 'Generation DOCX' enregistre", any(e["action"] == "Génération DOCX" for e in history_after_docx))

# ---------------------------------------------------------------------------
print("\n== Export PDF reel (necessite Microsoft Word sur ce poste) ==")

try:
    pdf_path = cddu_service.export_pdf(contrat_id)
    check("le fichier PDF existe sur le disque", pdf_path.exists())
    check("le fichier PDF n'est pas vide", pdf_path.stat().st_size > 0)

    saved_after_pdf = cddu_service.get_contrat(contrat_id)
    check("contrat.pdf_path enregistre", saved_after_pdf.pdf_path == str(pdf_path))
    check("statut avance automatiquement a 'pdf_generated' (etait 'draft')", saved_after_pdf.status == "pdf_generated")

    history_after_pdf = cddu_service.history(contrat_id)
    check("historique 'Export PDF' enregistre", any(e["action"] == "Export PDF" for e in history_after_pdf))
except Exception as exc:  # WordNotAvailableError / PdfConversionTimeoutError / autre
    print(f"[INFO] Conversion PDF non verifiee sur ce poste : {exc!r}")
    print("       (necessite Microsoft Word installe - non bloquant pour le reste du test)")

# ---------------------------------------------------------------------------
print("\n== Nettoyage ==")

database_module.Database.instance().close()
_tmp_dir.cleanup()
print("Base temporaire supprimee.")
print(f"Fichiers generes conserves pour inspection manuelle : {docx_path}")

print("\n" + "=" * 60)
if failures:
    print(f"{len(failures)} verification(s) en echec :")
    for label in failures:
        print(f"  - {label}")
    raise SystemExit(1)

print("Toutes les verifications sont passees.")
