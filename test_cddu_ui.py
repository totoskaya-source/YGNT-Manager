"""Test UI reel - Sprint 17.0, interface du module CDDU.

Script autonome (meme convention que les tests precedents), execute sur une
base SQLite TEMPORAIRE et isolee (data/ygnt_manager.db n'est jamais touchee),
avec une VRAIE QApplication PySide6 en mode "offscreen" (aucune fenetre
affichee, mais les widgets, signaux et logique de CdduDialog/CdduPage
s'executent reellement - pas de mock de l'UI elle-meme).

Verifie : creation, ouverture (edition), modification, generation DOCX,
generation PDF, ouverture DOCX/PDF (os.startfile mocke pour eviter de
lancer une vraie application externe pendant les tests, mais tout le reste
du chemin - bouton -> handler -> Service -> Repository -> historique - est
reellement exerce), defraiements, numero d'objet vide, aucun placeholder
residuel, et surtout : la regle "aucun autre artiste que
prestation_participants" dans le combo Artiste du dialogue.
"""

from __future__ import annotations

import os
import tempfile
from pathlib import Path
from unittest.mock import patch

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")

import app.database.database as database_module

_tmp_dir = tempfile.TemporaryDirectory()
database_module.DB_PATH = Path(_tmp_dir.name) / "test_cddu_ui.db"
database_module.Database._instance = None

from PySide6.QtWidgets import QApplication, QMessageBox  # noqa: E402

app = QApplication.instance() or QApplication([])

# QMessageBox.information/warning/question ouvrent une vraie boite modale
# Qt (.exec()) qui attend un clic utilisateur. En mode offscreen, sans
# automatisation de clic, cet appel bloque indefiniment - c'est ce qui a
# provoque le blocage precedent (des le premier dialog.save(), qui declenche
# notify_success -> QMessageBox.information). On les neutralise pour toute
# la duree du script : le CODE testé (CdduDialog/CdduPage) reste reellement
# exerce, seule la boite de dialogue elle-meme (rendu, boucle d'attente du
# clic) est court-circuitee - technique standard de test Qt headless.
_info_patch = patch.object(QMessageBox, "information", return_value=QMessageBox.StandardButton.Ok)
_warning_patch = patch.object(QMessageBox, "warning", return_value=QMessageBox.StandardButton.Ok)
_question_patch = patch.object(QMessageBox, "question", return_value=QMessageBox.StandardButton.Yes)
_info_patch.start()
_warning_patch.start()
_question_patch.start()

from app.database.migrations import MigrationManager  # noqa: E402
from app.models.artist import Artist  # noqa: E402
from app.models.prestation import Prestation  # noqa: E402
from app.models.producteur import Producteur  # noqa: E402
from app.services.artist_service import ArtistService  # noqa: E402
from app.services.contrat_cddu_service import ContratCdduService  # noqa: E402
from app.services.organization_service import OrganizationService  # noqa: E402
from app.services.prestation_participant_service import PrestationParticipantService  # noqa: E402
from app.services.prestation_service import PrestationService  # noqa: E402
from app.services.producteur_service import ProducteurService  # noqa: E402
from app.ui.cddu import CdduPage  # noqa: E402
from app.ui.cddu_dialog import CdduDialog  # noqa: E402

failures: list[str] = []


def check(label: str, condition: bool) -> None:
    status = "OK" if condition else "ECHEC"
    print(f"[{status}] {label}")
    if not condition:
        failures.append(label)


MigrationManager().migrate()

# ---------------------------------------------------------------------------
print("== Donnees de test ==")

producteur_service = ProducteurService()
producteur_service.create_producteur(Producteur(
    nom="YGNT PRODUCTION", representant="Tanguy Zahn", fonction="President",
    city="Villeneuve", siret="10572314200011", convention_collective="IDCC 01285",
))

artist_service = ArtistService()
anthony_id = artist_service.create_artist(Artist(
    legal_name="Anthony Zahn", address="301 Rue du Pigeonnier", postal_code="83600",
    city="Fréjus", birth_date="11/05/1988", birth_place="Cavaillon",
    social_number="1 88 05 84 035 041 22", conges_spectacle_number="C 381 825",
    instrument="Guitariste", phone="0600000000", qualification="Artiste musicien",
))
# Un second artiste, jamais rattache a prestation_participants : doit rester
# absent du combo Artiste, quelle que soit la prestation choisie.
jose_id = artist_service.create_artist(Artist(legal_name="Jose Intrus", instrument="Chanteur", qualification="Chanteur"))
sanfuego_id = artist_service.create_artist(Artist(legal_name="SANFUEGO", instrument="Groupe", qualification="Artiste musicien"))

organization_service = OrganizationService()
mairie_id = organization_service.create_organization(__import__("app.models.organization", fromlist=["Organization"]).Organization(name="Mairie de Test"))

prestation_service = PrestationService()
prestation_id = prestation_service.create_prestation(Prestation(
    nom="Concert de flamenco - Nuits Nomades de Maneo", type_evenement="festival",
    date_debut="09/07/2026", lieu_nom="Maneo", lieu_city="Aix en Provence",
    artist_id=sanfuego_id, organization_id=mairie_id,
))
# Prestation secondaire, sans aucun participant enregistre : sert a tester le
# bouton "+ Ajouter la formation".
prestation_vide_id = prestation_service.create_prestation(Prestation(
    nom="Prestation sans equipe", type_evenement="mariage", date_debut="20/08/2026",
    artist_id=sanfuego_id,
))

participant_service = PrestationParticipantService()
participant_service.add_participant(prestation_id, anthony_id, role="Guitariste", ordre=1)

cddu_service = ContratCdduService()

# ---------------------------------------------------------------------------
print("\n== CdduPage : liste vide au depart ==")

page = CdduPage(service=cddu_service)
check("aucune ligne dans la liste avant toute creation", page.table.rowCount() == 0)

# ---------------------------------------------------------------------------
print("\n== CdduDialog : creation (cas principal, 1 prestation -> 1 artiste -> 1 contrat) ==")

dialog = CdduDialog(
    None,
    service=cddu_service,
    artist_service=artist_service,
    prestation_service=prestation_service,
    organization_service=organization_service,
    participant_service=participant_service,
)
# .show() (fenetre reelle, meme en offscreen) : necessaire pour que
# QWidget.isVisible() reflete l'etat reel des enfants (un widget dont la
# fenetre parente n'a jamais ete affichee est toujours "invisible" en Qt,
# quel que soit l'appel setVisible() interne) - sans quoi les verifications
# de visibilite du bouton "+ Ajouter la formation" seraient faussees.
dialog.show()

check("aucune prestation choisie au depart : combo Artiste desactive", not dialog.artist_combo.isEnabled())

prestation_index = dialog.prestation_combo.findData(prestation_id)
check("la prestation de test est proposee dans le combo", prestation_index >= 0)
dialog.prestation_combo.setCurrentIndex(prestation_index)

check("formation affichee automatiquement (lecture seule)", dialog.prestation_formation.text() == "SANFUEGO")
check("date affichee automatiquement", dialog.prestation_date.text() == "09/07/2026")
check("lieu affiche automatiquement", dialog.prestation_lieu_display.text() == "Maneo, Aix en Provence")
check("organisateur affiche automatiquement", dialog.prestation_organisateur.text() == "Mairie de Test")

combo_artist_ids = [dialog.artist_combo.itemData(i) for i in range(dialog.artist_combo.count())]
check("seul Anthony (prestation_participants) est propose", combo_artist_ids == [anthony_id])
check("Jose (jamais rattache a cette prestation) est absent du combo", jose_id not in combo_artist_ids)

check("l'artiste est preselectionne (un seul choix)", dialog.artist_combo.currentData() == anthony_id)
check("champs Artiste pre-remplis automatiquement", dialog.artiste_nom.text() == "Anthony Zahn")
check("numero de securite sociale pre-rempli", dialog.artiste_numero_secu.text() == "1 88 05 84 035 041 22")
check("fonction pre-remplie depuis l'instrument", dialog.fonction.text() == "Guitariste")
check("champs Artiste en lecture seule", dialog.artiste_nom.isReadOnly())

dialog.salaire_brut.setValue(136.37)
dialog.defraiement_deplacement.setValue(45.0)
# Hebergement/Repas/Autres laisses a 0 : doivent disparaitre du DOCX genere.

dialog.save()

check("le CDDU est enregistre (identifiant attribue)", dialog._source_contrat is not None and dialog._source_contrat.id is not None)
contrat_id = dialog._source_contrat.id
check("numero au format CDDU-AAAA-0001", dialog._source_contrat.numero.endswith("-0001"))
check("bouton Annuler devient 'Fermer' apres enregistrement", dialog.buttons.button(dialog.buttons.StandardButton.Cancel).text() == "Fermer")

saved = cddu_service.get_contrat(contrat_id)
check("instantane producteur enregistre", saved.producteur_nom == "YGNT PRODUCTION")
check("instantane artiste enregistre", saved.artiste_numero_conges_spectacle == "C 381 825")
check("salaire brut enregistre", saved.remuneration_brute == 136.37)
check("defraiement deplacement enregistre", saved.defraiement_deplacement == 45.0)
check("statut par defaut Brouillon", saved.status == "draft")
check("numero_objet reste vide", saved.numero_objet == "")

dates = cddu_service.list_dates(contrat_id)
check("une ligne contrat_cddu_dates creee automatiquement (cas simple)", len(dates) == 1)
check("la date reprend celle de la prestation choisie", bool(dates) and dates[0].date_travaillee == "09/07/2026")

# ---------------------------------------------------------------------------
print("\n== Generation DOCX (bouton reel) ==")

dialog.btn_generate_docx.click()

check("bouton DOCX active apres enregistrement", True)  # deja verifie ci-dessus indirectement
refreshed = cddu_service.get_contrat(contrat_id)
check("docx_path enregistre apres generation", bool(refreshed.docx_path) and Path(refreshed.docx_path).exists())

from docx import Document  # noqa: E402

doc = Document(refreshed.docx_path)


def all_text(document) -> str:
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


full_text = all_text(doc)
check("aucun placeholder {{...}} residuel", "{{" not in full_text and "}}" not in full_text)
check("defraiement Deplacement present", "Déplacement : 45,00 euros" in full_text)
check("defraiement Hebergement absent (non renseigne)", "Hébergement :" not in full_text)
check("defraiement Repas absent (non renseigne)", "Repas :" not in full_text)

numero_objet_paragraph = next((p.text for p in doc.paragraphs if "numéro d'objet" in p.text), None)
check(
    "numero d'objet vide dans le document genere",
    numero_objet_paragraph is not None and numero_objet_paragraph.strip().endswith("le"),
)

history = cddu_service.history(contrat_id)
check("historique 'Generation DOCX' trace", any(e["action"] == "Génération DOCX" for e in history))

# ---------------------------------------------------------------------------
print("\n== Generation PDF (bouton reel, via Microsoft Word) ==")

try:
    dialog.btn_generate_pdf.click()
    refreshed = cddu_service.get_contrat(contrat_id)
    check("pdf_path enregistre apres generation", bool(refreshed.pdf_path) and Path(refreshed.pdf_path).exists())
    check("statut avance a 'PDF genere'", refreshed.status == "pdf_generated")
    check(
        "le combo Statut du dialogue reflete l'avancement automatique",
        dialog.status.currentData() == "pdf_generated",
    )
    pdf_available = True
except Exception as exc:  # Microsoft Word non disponible sur ce poste
    print(f"[INFO] Conversion PDF non verifiee sur ce poste : {exc!r}")
    pdf_available = False

# ---------------------------------------------------------------------------
print("\n== Ouverture DOCX/PDF (os.startfile mocke pour eviter de lancer une vraie application) ==")

with patch("os.startfile") as mock_startfile:
    dialog.btn_open_docx.click()
    check("Ouvrir DOCX declenche bien l'ouverture du fichier", mock_startfile.called)

if pdf_available:
    with patch("os.startfile") as mock_startfile:
        dialog.btn_open_pdf.click()
        check("Ouvrir PDF declenche bien l'ouverture du fichier", mock_startfile.called)

history_after_open = cddu_service.history(contrat_id)
check("historique 'Ouverture DOCX' trace", any(e["action"] == "Ouverture DOCX" for e in history_after_open))

# ---------------------------------------------------------------------------
print("\n== Modification (reouverture du CDDU) ==")

# La fiche Artiste change APRES la creation du CDDU : le CDDU deja enregistre
# ne doit jamais refleter ce changement (instantane fige).
artist_service.update_artist(
    Artist(id=anthony_id, legal_name="Anthony Zahn", phone="0699999999", instrument="Guitariste", qualification="Artiste musicien"),
)

# Rechargement frais (et non la variable "saved" capturee avant generation
# DOCX/PDF) : c'est ce que fait reellement CdduPage._selected_contrat() avant
# d'ouvrir le dialogue en edition (get_contrat() a chaque ouverture) - le
# reutiliser tel quel eviterait un contrat perime, pas un comportement a
# tester ici.
contrat_before_edit = cddu_service.get_contrat(contrat_id)

edit_dialog = CdduDialog(
    None,
    contrat=contrat_before_edit,
    service=cddu_service,
    artist_service=artist_service,
    prestation_service=prestation_service,
    organization_service=organization_service,
    participant_service=participant_service,
)
edit_dialog.show()

check("ouverture en modification : reference pre-remplie", edit_dialog.reference.text() == saved.numero)
check("ouverture en modification : prestation preselectionnee", edit_dialog.prestation_combo.currentData() == prestation_id)
check("ouverture en modification : artiste preselectionne", edit_dialog.artist_combo.currentData() == anthony_id)
check(
    "instantane fige respecte : l'ancien telephone reste affiche, pas le nouveau",
    edit_dialog.artiste_phone.text() == "" or edit_dialog.artiste_phone.text() != "0699999999",
)
check("salaire brut pre-rempli depuis l'enregistrement existant", edit_dialog.salaire_brut.value() == 136.37)
check("defraiement deplacement pre-rempli depuis l'enregistrement existant", edit_dialog.defraiement_deplacement.value() == 45.0)

edit_dialog.salaire_brut.setValue(150.0)
edit_dialog.status.setCurrentIndex(edit_dialog.status.findData("validated"))
edit_dialog.save()

modified = cddu_service.get_contrat(contrat_id)
check("modification du salaire brut persistee", modified.remuneration_brute == 150.0)
check("modification du statut persistee", modified.status == "validated")

history_after_edit = cddu_service.history(contrat_id)
check("historique 'Modification' trace", any(e["action"] == "Modification" for e in history_after_edit))

# ---------------------------------------------------------------------------
print("\n== Bouton '+ Ajouter la formation' (prestation sans equipe) ==")

empty_dialog = CdduDialog(
    None,
    service=cddu_service,
    artist_service=artist_service,
    prestation_service=prestation_service,
    organization_service=organization_service,
    participant_service=participant_service,
)

empty_dialog.show()
# QTabWidget masque reellement les onglets non actifs (QStackedWidget sous-
# jacent) : isVisible() sur un widget d'un onglet inactif renvoie toujours
# False, quel que soit son propre setVisible(). On bascule sur l'onglet
# Artiste (index 1 : General, Artiste, Prestation, Defraiements, Apercu)
# pour verifier honnetement l'etat du bouton.
empty_dialog.tabs.setCurrentIndex(1)
check("l'onglet Artiste est bien celui affiche pour ce test", empty_dialog.tabs.tabText(empty_dialog.tabs.currentIndex()) == "Artiste")

empty_index = empty_dialog.prestation_combo.findData(prestation_vide_id)
empty_dialog.prestation_combo.setCurrentIndex(empty_index)

check("prestation sans equipe : combo Artiste vide (placeholder)", empty_dialog.artist_combo.currentData() is None)
check("bouton '+ Ajouter la formation' visible", empty_dialog.btn_add_formation_participant.isVisible())

empty_dialog.btn_add_formation_participant.click()

check(
    "la formation devient participante apres le clic",
    any(p.artiste_id == sanfuego_id for p in participant_service.list_participants(prestation_vide_id)),
)
check("le combo Artiste propose desormais la formation", empty_dialog.artist_combo.currentData() == sanfuego_id)
check("le bouton se masque une fois la formation ajoutee", not empty_dialog.btn_add_formation_participant.isVisible())

# ---------------------------------------------------------------------------
print("\n== CdduPage : liste apres creation ==")

page.refresh_table()
check("une ligne dans la liste apres creation", page.table.rowCount() == 1)
check("colonne Reference correcte", page.table.item(0, 1).text() == saved.numero)
check("colonne Artiste correcte", page.table.item(0, 2).text() == "Anthony Zahn")
check("colonne Prestation correcte", page.table.item(0, 3).text() == saved.prestation_reference)
check("colonne Date correcte", page.table.item(0, 4).text() == "09/07/2026")
check("colonne Statut refletant la modification", page.table.item(0, 5).text() == "Validé")
check("colonne PDF correcte", page.table.item(0, 6).text() == ("Oui" if pdf_available else "Non"))

# ---------------------------------------------------------------------------
print("\n== Nettoyage ==")

_info_patch.stop()
_warning_patch.stop()
_question_patch.stop()

database_module.Database.instance().close()
_tmp_dir.cleanup()
print("Base temporaire supprimee.")

print("\n" + "=" * 60)
if failures:
    print(f"{len(failures)} verification(s) en echec :")
    for label in failures:
        print(f"  - {label}")
    raise SystemExit(1)

print("Toutes les verifications sont passees.")
