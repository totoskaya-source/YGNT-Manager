"""Test reel - Sprint 18.2, relecture redactionnelle du modele CDDU.

Genere 4 CDDU reels (musicien, chanteur, danseuse, technicien) et verifie que
la lecture est naturelle et juridiquement neutre : qualification variable
(jamais codee en dur), instrument/specialite optionnel (disparait si vide),
nom complet (prenom + NOM) partout, defraiements integralement absents si
rien n'est renseigne, en-tete complet, aucun placeholder residuel.

Base SQLite TEMPORAIRE et isolee (data/ygnt_manager.db n'est jamais touchee).
"""

from __future__ import annotations

import tempfile
from pathlib import Path

import app.database.database as database_module

_tmp_dir = tempfile.TemporaryDirectory()
database_module.DB_PATH = Path(_tmp_dir.name) / "test_cddu_wording_review.db"
database_module.Database._instance = None

from app.database.migrations import MigrationManager  # noqa: E402
from app.models.artist import Artist  # noqa: E402
from app.models.prestation import Prestation  # noqa: E402
from app.models.producteur import Producteur  # noqa: E402
from app.services.artist_service import ArtistService  # noqa: E402
from app.services.contrat_cddu_service import ContratCdduService  # noqa: E402
from app.services.prestation_service import PrestationService  # noqa: E402
from app.services.producteur_service import ProducteurService  # noqa: E402

failures: list[str] = []


def check(label: str, condition: bool) -> None:
    status = "OK" if condition else "ECHEC"
    print(f"[{status}] {label}")
    if not condition:
        failures.append(label)


def full_text_of(docx_path) -> str:
    from docx import Document

    doc = Document(str(docx_path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


MigrationManager().migrate()

# ---------------------------------------------------------------------------
print("== Donnees de test ==")

producteur_service = ProducteurService()
producteur_service.create_producteur(Producteur(
    nom="YGNT PRODUCTION", representant="Sophie Dupont", fonction="Présidente",
    city="Villeneuve", siret="10572314200011",
    convention_collective="Convention collective nationale des entreprises artistiques et culturelles (IDCC 1285)",
))

artist_service = ArtistService()
prestation_service = PrestationService()
cddu_service = ContratCdduService()

prestation_id = prestation_service.create_prestation(Prestation(
    nom="Festival Été 2026", type_evenement="festival", date_debut="09/07/2026",
    lieu_nom="Plaine des fêtes", lieu_city="Testville",
))
prestation = prestation_service.get_prestation(prestation_id)

PERSONAS = [
    {
        "cle": "musicien",
        "first_name": "Anthony",
        "legal_name": "Zahn",
        "instrument": "Guitare",
        "qualification": "Artiste musicien",
        "defraiement_deplacement": "45",
    },
    {
        "cle": "chanteur",
        "first_name": "Mahoni",
        "legal_name": "Ravel",
        "instrument": "Chant",
        "qualification": "Chanteur",
        "defraiement_deplacement": "0",
    },
    {
        "cle": "danseuse",
        "first_name": "Léa",
        "legal_name": "Martin",
        "instrument": "",  # une danseuse n'a pas d'instrument : ligne attendue absente
        "qualification": "Danseur",
        "defraiement_deplacement": "0",
    },
    {
        "cle": "technicien",
        "first_name": "Karim",
        "legal_name": "Belkacem",
        "instrument": "",  # un technicien n'a pas d'instrument : ligne attendue absente
        "qualification": "Technicien du spectacle",
        "defraiement_deplacement": "0",
    },
]

generated = {}

for persona in PERSONAS:
    artist_id = artist_service.create_artist(Artist(
        legal_name=persona["legal_name"],
        first_name=persona["first_name"],
        instrument=persona["instrument"],
        qualification=persona["qualification"],
    ))

    contrat = cddu_service.build_from_prestation_and_artist(prestation, artist_id)
    contrat.remuneration_brute = "150"
    contrat.defraiement_deplacement = persona["defraiement_deplacement"]

    contrat_id = cddu_service.create_contrat(contrat)
    cddu_service.add_date(contrat_id, "09/07/2026", prestation_id=prestation_id, nombre_cachets=1)

    docx_path = cddu_service.generate_docx(contrat_id)
    text = full_text_of(docx_path)
    generated[persona["cle"]] = {"id": contrat_id, "path": docx_path, "text": text, "persona": persona}

    print(f"\n== {persona['cle'].upper()} : {persona['first_name']} {persona['legal_name']} ==")

    check(f"[{persona['cle']}] aucun placeholder residuel", "{{" not in text and "}}" not in text)

    nom_complet = f"{persona['first_name']} {persona['legal_name'].upper()}"
    check(f"[{persona['cle']}] nom complet 'Prenom NOM' present ({nom_complet})", nom_complet in text)
    check(f"[{persona['cle']}] le seul nom de famille n'apparait jamais isole", persona["legal_name"] not in text.replace(nom_complet, ""))

    check(
        f"[{persona['cle']}] qualification correcte affichee (en qualité de {persona['qualification']})",
        f"en qualité de {persona['qualification']}." in text,
    )
    check(f"[{persona['cle']}] jamais 'artiste musicien' code en dur pour un non-musicien" if persona["qualification"] != "Artiste musicien" else f"[{persona['cle']}] mention musicien coherente", (
        "en qualité d'artiste musicien" not in text.lower() if persona["qualification"] != "Artiste musicien" else True
    ))

    if persona["instrument"]:
        check(f"[{persona['cle']}] ligne Instrument principal presente", f"Instrument principal :\n{persona['instrument']}" in text)
    else:
        check(f"[{persona['cle']}] ligne Instrument principal absente (aucun instrument renseigne)", "Instrument principal" not in text)

    if float(persona["defraiement_deplacement"]) == 0:
        check(f"[{persona['cle']}] article DEFRAIEMENTS entierement absent (titre compris)", "DÉFRAIEMENTS" not in text)
    else:
        check(f"[{persona['cle']}] article DEFRAIEMENTS present avec le bon montant", "DÉFRAIEMENTS" in text and "Déplacement : 45,00 euros" in text)
        check(f"[{persona['cle']}] les lignes non renseignees restent absentes", "Hébergement :" not in text and "Repas :" not in text)

    check(f"[{persona['cle']}] en-tete : numero present", generated[persona["cle"]]["id"] and text.split("\n")[1].startswith("Contrat CDDU-"))
    check(f"[{persona['cle']}] en-tete : prestation et date presentes", "Festival Été 2026" in text and "09/07/2026" in text)

# ---------------------------------------------------------------------------
print("\n== Signatures : colonnes uniformes ==")

from docx import Document  # noqa: E402

sample_doc = Document(str(generated["musicien"]["path"]))
check("le tableau de signatures existe", len(sample_doc.tables) == 1)
signature_table = sample_doc.tables[0]
widths = [cell.width for cell in signature_table.rows[0].cells]
check("les deux colonnes de signature ont exactement la meme largeur", widths[0] == widths[1] and widths[0] is not None)

alignments = [cell.paragraphs[0].alignment for row in signature_table.rows for cell in row.cells]
check("toutes les cellules du tableau de signature partagent le meme alignement", len(set(alignments)) == 1)

# ---------------------------------------------------------------------------
print("\n== PDF reel (musicien, verifie que la conversion fonctionne toujours) ==")

try:
    pdf_path = cddu_service.export_pdf(generated["musicien"]["id"])
    check("le PDF se genere toujours correctement apres la relecture redactionnelle", pdf_path.exists())

    import fitz  # noqa: E402

    pdf_text = "".join(page.get_text() for page in fitz.open(str(pdf_path)))
    check("le PDF reprend le nom complet Anthony ZAHN", "Anthony ZAHN" in pdf_text)
    check("le PDF reprend la qualification 'Artiste musicien'", "en qualité de Artiste musicien" in pdf_text)
except Exception as exc:
    print(f"[INFO] Conversion PDF non verifiee sur ce poste : {exc!r}")

# ---------------------------------------------------------------------------
print("\n== Non-regression : le CDDU 'musicien' se lit integralement ==")

musicien_text = generated["musicien"]["text"]
for expected_heading in ("OBJET", "DURÉE DE L'ENGAGEMENT", "RÉMUNÉRATION", "DÉFRAIEMENTS", "RUPTURE ANTICIPÉE POUR FAUTE GRAVE OU FORCE MAJEURE", "RETRAITE ET CONGÉS PAYÉS", "ABSENCE - MALADIE", "MÉDECINE DU TRAVAIL", "ASSURANCES", "LITIGES"):
    check(f"article '{expected_heading}' toujours present", expected_heading in musicien_text)

# ---------------------------------------------------------------------------
print("\n== Nettoyage ==")

database_module.Database.instance().close()
_tmp_dir.cleanup()
print("Base temporaire supprimee.")

print("\n" + "=" * 60)
if failures:
    print(f"{len(failures)} verification(s) en echec :")
    for label in failures:
        print(f"  - {label}")
    raise SystemExit(1)

print("Toutes les verifications sont passees.")
