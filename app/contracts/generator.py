from pathlib import Path

from app.documents.placeholder_engine import PlaceholderEngine


class ContractGenerator:
    ORGANIZER_DETAILS = (
        ("Forme juridique", "organisateur_forme"),
        ("Adresse", "organisateur_adresse"),
        ("Code postal", "organisateur_postal_code"),
        ("Ville", "organisateur_city"),
        ("SIRET", "organisateur_siret"),
        ("Telephone", "organisateur_phone"),
        ("Email", "organisateur_email"),
        ("Code APE", "organisateur_ape"),
        ("Licence spectacle", "organisateur_licence"),
        ("TVA intracommunautaire", "organisateur_tva"),
        ("Representee par", "organisateur_representant"),
        ("Fonction du representant", "organisateur_fonction"),
    )

    def __init__(self, template_path):
        self.template_path = Path(template_path)

    def generate(self, contract, output_path):
        from docx import Document

        doc = Document(self.template_path)
        values = contract.to_dict()

        has_structure_indicator = self._has_structure_indicator(values)
        values["organisateur_structure_label"] = (
            "Structure" if has_structure_indicator else "Nom et prénom"
        )
        if not has_structure_indicator:
            # Un particulier n'a ni SIRET, ni representant, ni fonction, ni
            # TVA, ni licence : ces champs ne doivent jamais apparaitre pour
            # lui, meme si une valeur residuelle traine dans les donnees.
            for key in self.PARTICULIER_HIDDEN_KEYS:
                values[key] = ""

        self.replace_in_document(doc, values)

        doc.save(output_path)

    # Champs qui ne peuvent exister que pour une structure enregistree
    # (raison sociale, association, collectivite...) : un particulier n'a
    # ni forme juridique, ni SIRET, ni TVA intracommunautaire, ni licence
    # d'entrepreneur de spectacles. La presence d'un seul de ces champs
    # suffit a prouver qu'il ne s'agit pas d'une personne physique, meme si
    # la "Forme juridique" elle-meme n'a pas ete saisie.
    STRUCTURE_INDICATOR_KEYS = (
        "organisateur_forme",
        "organisateur_siret",
        "organisateur_tva",
        "organisateur_licence",
    )

    # Champs a toujours masquer pour un particulier (voir generate()).
    # Inclut STRUCTURE_INDICATOR_KEYS (deja garantis vides par definition)
    # ainsi que representant/fonction, qui n'en font pas partie.
    PARTICULIER_HIDDEN_KEYS = STRUCTURE_INDICATOR_KEYS + (
        "organisateur_representant",
        "organisateur_fonction",
    )

    @staticmethod
    def _has_structure_indicator(values):
        """Un organisateur particulier n'a pas de forme juridique (SARL,
        association, mairie...) : dans ce cas le champ nom affiche une
        personne physique, pas une structure."""
        return any(
            str(values.get(key) or "").strip()
            for key in ContractGenerator.STRUCTURE_INDICATOR_KEYS
        )

    def replace_in_document(self, doc, values):
        PlaceholderEngine.replace_in_document(doc, values)

    def replace_paragraph(self, paragraph, values):
        PlaceholderEngine.replace_paragraph(paragraph, values)

    def append_organizer_details(self, doc, values):
        details = [
            (label, values.get(key, ""))
            for label, key in self.ORGANIZER_DETAILS
            if values.get(key)
        ]

        if not details:
            return

        doc.add_paragraph()
        title = doc.add_paragraph("Informations organisateur")
        title.runs[0].bold = True

        table = doc.add_table(rows=0, cols=2)

        for label, value in details:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = str(value)
