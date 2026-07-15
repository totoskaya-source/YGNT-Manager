from __future__ import annotations

from pathlib import Path
from typing import Iterable

from app.documents.placeholder_engine import PlaceholderEngine
from app.models.contrat_cddu_date import ContratCdduDate


class CdduGenerator:
    """Genere le DOCX d'un Contrat de travail (CDDU) - meme modele que
    ContractGenerator (app/contracts/generator.py) : ouvre le template,
    calcule les valeurs, delegue la substitution a PlaceholderEngine, sans
    aucune duplication de sa logique."""

    # Montants de defraiement optionnels : une valeur nulle/vide doit
    # disparaitre du document (jamais "0,00 euros"), voir _blank_zero_montants.
    # PlaceholderEngine ne retire un paragraphe que si la valeur brute est
    # une chaine vide/None, pas un nombre a zero (voir sa docstring).
    DEFRAIEMENT_MONTANT_KEYS = (
        "defraiement_deplacement",
        "defraiement_hebergement",
        "defraiement_repas",
        "defraiement_autres_montant",
        "defraiement_montant_libre_montant",
    )

    def __init__(self, template_path):
        self.template_path = Path(template_path)

    # Titre d'article a retirer entierement (pas seulement ses lignes) si
    # aucun defraiement n'est renseigne (Sprint 18.2 §3) - un paragraphe de
    # titre ne contient aucun placeholder, donc la regle generique de
    # PlaceholderEngine ("supprime un paragraphe si tous ses placeholders
    # sont vides") ne s'applique jamais a lui : traite explicitement ici.
    DEFRAIEMENTS_HEADING = "DÉFRAIEMENTS"

    def generate(self, contrat, dates: Iterable[ContratCdduDate], output_path) -> None:
        from docx import Document

        doc = Document(self.template_path)

        values = contrat.to_dict()
        values.update(self._date_values(dates))
        self._blank_zero_montants(values)

        # BUG-002 (v1.0.3) : producteur_adresse_complete/artiste_adresse_complete
        # (calculees par ContratCddu.to_dict()) ignorent deja un champ
        # postal_code/city vide, mais pas une virgule parasite deja presente
        # dans le champ Adresse lui-meme (simple texte libre, ex. saisi avec
        # une virgule de fin) : "adresse," + ", CP VILLE" => double virgule
        # a l'affichage. Reconstruites ici a partir des memes champs bruts,
        # chacun nettoye de ses virgules/espaces parasites avant assemblage -
        # le format final (une seule virgule entre adresse et "CP VILLE")
        # est inchange quand les donnees sont propres.
        values["producteur_adresse_complete"] = self._build_adresse_complete(
            values.get("producteur_adresse"),
            values.get("producteur_postal_code"),
            values.get("producteur_city"),
        )
        values["artiste_adresse_complete"] = self._build_adresse_complete(
            values.get("artiste_adresse"),
            values.get("artiste_postal_code"),
            values.get("artiste_city"),
        )

        # Alias de lecture (Sprint 18.2) : le contrat n'introduit aucune
        # nouvelle colonne pour ces deux notions, elles reutilisent des
        # champs deja existants sous un nom de placeholder plus clair pour
        # le document imprime. Aucune valeur n'est jamais codee en dur ici :
        # {{qualification}} et {{instrument_principal}} restent vides si la
        # fiche Artiste ne les a pas renseignes.
        values["qualification"] = values.get("artiste_qualification", "")
        values["instrument_principal"] = values.get("artiste_fonction", "")

        # Clause de fonction du representant (point orphelin) : le template
        # ecrit "represente par {{producteur_representant}}{{producteur_fonction_clause}}."
        # - la clause complete (", agissant en qualite de X") n'apparait que
        # si producteur_fonction est renseignee, sinon elle est une chaine
        # vide et la phrase se termine directement apres le nom, sans virgule
        # ni point orphelin.
        producteur_fonction = str(values.get("producteur_fonction") or "").strip()
        values["producteur_fonction_clause"] = (
            f", agissant en qualité de {producteur_fonction}" if producteur_fonction else ""
        )

        # {{numero_objet}} doit rester visible meme vide (la phrase du
        # contrat de reference reste presente, juste sans numero apres "le"),
        # voir docs/CDDU_ARCHITECTURE.md §9. Traite avant/hors du passage
        # generique : PlaceholderEngine supprimerait sinon le paragraphe
        # entier, puisque numero_objet est le seul placeholder de sa ligne
        # et une valeur vide declenche sa regle de suppression des
        # paragraphes entierement vides.
        numero_objet = values.pop("numero_objet", "")
        self._apply_numero_objet(doc, numero_objet)

        self._remove_defraiements_heading_if_empty(doc, values)

        self.replace_in_document(doc, values)
        self._expand_line_breaks(doc)

        doc.save(output_path)

    def replace_in_document(self, doc, values):
        PlaceholderEngine.replace_in_document(doc, values)

    @staticmethod
    def _clean_address_part(value) -> str:
        """Normalise un fragment d'adresse avant assemblage (BUG-002,
        v1.0.3) : espaces superflus et virgule(s) de fin retirees, pour
        qu'un champ deja termine par une virgule (saisie manuelle) ne
        produise jamais de double separateur une fois combine aux autres
        fragments. Une chaine vide ou None reste une chaine vide, exclue
        normalement de l'assemblage (voir _build_adresse_complete)."""
        text = str(value or "").strip()
        return text.rstrip(",").strip()

    @classmethod
    def _build_adresse_complete(cls, adresse, postal_code, city) -> str:
        """Assemble "adresse, CP VILLE" en ignorant tout fragment vide -
        jamais de virgule ni d'espace parasite si l'un des trois manque
        (meme format final que ContratCddu.producteur_adresse_complete /
        artiste_adresse_complete, dont cette methode reprend le resultat
        apres nettoyage des fragments, voir BUG-002)."""
        code_ville = " ".join(
            part for part in (cls._clean_address_part(postal_code), cls._clean_address_part(city)) if part
        )
        parts = (cls._clean_address_part(adresse), code_ville)
        return ", ".join(part for part in parts if part)

    @staticmethod
    def _date_values(dates: Iterable[ContratCdduDate]) -> dict:
        """Le contrat ne lit jamais directement ses propres colonnes de
        dates (il n'en a pas, voir docs/CDDU_ARCHITECTURE.md §3) : toutes
        les dates proviennent de contrat_cddu_dates, fournies ici par le
        Service. Une date par ligne, triees chronologiquement."""
        ordered = sorted(dates, key=lambda date: date.date_travaillee)

        lines = [
            f"{date.date_travaillee} - {date.nombre_cachets} "
            f"Cachet{'s' if date.nombre_cachets != 1 else ''} de représentation"
            for date in ordered
        ]

        return {
            "date_debut": ordered[0].date_travaillee if ordered else "",
            "date_fin": ordered[-1].date_travaillee if ordered else "",
            "dates_travaillees": "\n".join(lines),
            "nombre_total_cachets": sum(date.nombre_cachets for date in ordered),
        }

    @staticmethod
    def _apply_numero_objet(doc, numero_objet) -> None:
        token = "{{numero_objet}}"
        replacement = str(numero_objet or "")

        for paragraph in doc.paragraphs:
            if not paragraph.runs:
                continue
            full_text = "".join(run.text for run in paragraph.runs)
            if token not in full_text:
                continue

            full_text = full_text.replace(token, replacement)
            paragraph.runs[0].text = full_text
            for run in paragraph.runs[1:]:
                run.text = ""

    @classmethod
    def _blank_zero_montants(cls, values: dict) -> None:
        for key in cls.DEFRAIEMENT_MONTANT_KEYS:
            try:
                is_zero = float(values.get(key) or 0) == 0
            except (TypeError, ValueError):
                is_zero = False
            if is_zero:
                values[key] = ""

    @classmethod
    def _remove_defraiements_heading_if_empty(cls, doc, values: dict) -> None:
        """Supprime l'article Defraiements dans son ensemble (titre compris)
        si rien n'est renseigne - jamais un titre orphelin sans contenu en
        dessous (Sprint 18.2 §3). Les lignes elles-memes disparaissent deja
        automatiquement via _blank_zero_montants + la regle generique de
        PlaceholderEngine ; seul le titre, sans placeholder, echappe a
        cette regle et doit etre traite explicitement."""
        has_content = any(values.get(key) for key in cls.DEFRAIEMENT_MONTANT_KEYS)
        has_content = has_content or bool(
            str(values.get("defraiement_autres_libelle") or "").strip()
            or str(values.get("defraiement_montant_libre_libelle") or "").strip()
        )
        if has_content:
            return

        for paragraph in doc.paragraphs:
            if paragraph.text.strip() == cls.DEFRAIEMENTS_HEADING:
                PlaceholderEngine._remove_paragraph(paragraph)
                return

    @staticmethod
    def _expand_line_breaks(doc) -> None:
        """Convertit les '\\n' internes issus d'une substitution multi-lignes
        (dates_travaillees) en veritables sauts de ligne Word. Reste local a
        ce generateur : PlaceholderEngine demeure generique et ignore cette
        mise en forme (voir sa docstring), aucune modification du moteur
        partage."""
        for paragraph in doc.paragraphs:
            for run in list(paragraph.runs):
                if "\n" not in run.text:
                    continue

                lines = run.text.split("\n")
                run.text = ""
                for index, line in enumerate(lines):
                    if index > 0:
                        run.add_break()
                    if line:
                        run.add_text(line)
