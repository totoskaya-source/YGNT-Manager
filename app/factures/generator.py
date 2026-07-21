from pathlib import Path

from app.documents.placeholder_engine import PlaceholderEngine

LOGO_MARKER = "{{producteur_logo}}"
LOGO_WIDTH_INCHES = 1.5


class FactureGenerator:
    def __init__(self, template_path):
        self.template_path = Path(template_path)

    def generate(self, facture, output_path):
        from docx import Document

        doc = Document(self.template_path)
        values = facture.to_dict()
        values.update(self._compute_ttc_values(facture))

        PlaceholderEngine.replace_in_document(doc, values)
        self._insert_logo(doc, getattr(facture, "producteur_logo_path", ""))

        doc.save(output_path)

    @staticmethod
    def _compute_ttc_values(facture):
        """Le Total TTC doit toujours etre affiche sur la facture, meme
        lorsqu'il est identique au Total HT (association non assujettie a la
        TVA). Facture.tva reste un champ texte libre non modifie ici (ex.
        "20%", "Exoneree") : un contenu non interpretable comme un taux
        numerique est traite comme 0%, de sorte que le TTC soit alors
        strictement egal au HT plutot que de laisser un total absent ou
        incoherent. Ces valeurs ne sont calculees que pour le contexte du
        document genere - Facture.tva et le reste du modele ne sont pas
        modifies."""
        total_ht = facture.montant
        taux_tva = FactureGenerator._parse_taux_tva(facture.tva)
        montant_tva = round(total_ht * taux_tva / 100, 2)
        total_ttc = round(total_ht + montant_tva, 2)

        return {
            "total_ht": total_ht,
            "tva": montant_tva,
            "taux_tva_label": FactureGenerator._format_taux(taux_tva),
            "total_ttc": total_ttc,
        }

    @staticmethod
    def _parse_taux_tva(tva_text):
        """Interprete le champ texte libre Facture.tva comme un taux en %.
        Une mention non numerique (vide, "Exoneree", "Non applicable"...)
        vaut 0%, jamais une erreur : la saisie utilisateur reste libre."""
        if not tva_text:
            return 0.0
        texte = str(tva_text).strip().replace("%", "").replace(",", ".")
        try:
            return float(texte)
        except ValueError:
            return 0.0

    @staticmethod
    def _format_taux(taux):
        """Formate un taux pour l'affichage ("TVA (5,5 %)") : entier sans
        decimale inutile, virgule francaise sinon."""
        texte = f"{taux:.2f}".rstrip("0").rstrip(".")
        return texte.replace(".", ",") or "0"

    def _insert_logo(self, doc, logo_path):
        """Remplace le marqueur {{producteur_logo}} par l'image du Producteur si
        elle existe sur le disque, sinon efface simplement le marqueur. Le
        chemin provient exclusivement de l'instantane stocke dans la Facture,
        jamais d'une relecture de la table Producteurs."""
        from docx.shared import Inches

        for paragraph in doc.paragraphs:
            if LOGO_MARKER not in paragraph.text:
                continue

            for run in paragraph.runs:
                run.text = ""

            if paragraph.runs:
                target_run = paragraph.runs[0]
            else:
                target_run = paragraph.add_run()

            if logo_path and Path(logo_path).exists():
                target_run.add_picture(logo_path, width=Inches(LOGO_WIDTH_INCHES))
